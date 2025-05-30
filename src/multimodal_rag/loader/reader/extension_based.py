import asyncio
from pathlib import Path
import json
import mimetypes
import chardet
import hashlib
from uuid import uuid4

from langdetect import detect, LangDetectException

from multimodal_rag.document import Document, SourceConfig, MetaConfig
from multimodal_rag.loader.reader.types import FileReader
from multimodal_rag.preprocessor.captioner.types import ImageCaptioner
from multimodal_rag.preprocessor.transcriber.types import AudioTranscriber
from multimodal_rag.log_config import logger
from multimodal_rag.utils.loader import load_image_bytes_from_source, load_file_from_path
from multimodal_rag.utils.timing import log_duration

LANG_EXT = {
    ".py": "python",
    ".js": "js",
    ".ts": "ts",
    ".java": "java",
    ".cpp": "cpp",
    ".c": "c",
    ".cs": "csharp",
    ".go": "go",
    ".php": "php",
    ".proto": "proto",
    ".kt": "kotlin",
    ".rb": "ruby",
    ".rs": "rust",
    ".scala": "scala",
    ".swift": "swift",
    ".sol": "sol",
    ".lua": "lua",
    ".pl": "perl",
    ".hs": "haskell",
    ".ex": "elixir",
    ".ps1": "powershell",
    ".tex": "latex",
    ".rst": "rst",
    ".md": "markdown",
    ".html": "html",
    ".cob": "cobol",
}


class ExtensionBasedReader(FileReader):
    """
    Reads a file based on its extension and MIME type.

    Supports code, text, markdown, HTML, PDF, DOCX, images, and audio.
    Uses captioner/transcriber for media if provided.
    """

    def __init__(self, transcriber: AudioTranscriber | None = None, captioner: ImageCaptioner | None = None):
        self.transcriber = transcriber
        self.captioner = captioner

    async def load(self, path: Path) -> list[Document]:
        path_str = str(path)
        ext = path.suffix.lower()
        mime, _ = mimetypes.guess_type(path_str)
        mime = mime or "application/octet-stream"

        logger.debug("Reading file", extra={"path": path_str, "ext": ext, "mime": mime})

        if ext in LANG_EXT:
            content = await self._read_text(path_str)
            content_type = f"code_{LANG_EXT[ext]}"
        elif ext == ".json":
            content = await self._read_json(path_str)
            content_type = "json"
        elif ext in {".txt", "", ".csv"}:
            content = await self._read_text(path_str)
            content_type = "text"
        elif ext == ".md":
            content = await self._read_text(path_str)
            content_type = "markdown"
        elif ext == ".html":
            content = await self._read_html(path_str)
            content_type = "markdown"
        elif ext == ".pdf":
            content = await self._read_pdf(path_str)
            content_type = "text"
        elif ext == ".docx":
            content = await self._read_docx(path_str)
            content_type = "text"
        else:
            if mime.startswith("image/"):
                content = await self._caption_image(path_str)
                content_type = "image"
            elif mime.startswith("audio/"):
                content = await self._transcribe_audio(path_str, mime)
                content_type = "text"
            else:
                content = ""
                content_type = "blob"

        try:
            lang = detect(content) if content else ""
        except LangDetectException:
            lang = ""
            logger.debug("Language detection failed", extra={"path": path_str})

        size_bytes = path.stat().st_size
        last_modified = int(path.stat().st_mtime)
        fingerprint = await self._hash_file(path)

        source_config = SourceConfig(
            tmp_path=path_str,
            loader="extension_based",
            type=content_type,
        )

        meta_config = MetaConfig(
            mime=mime,
            filename=path.name,
            size_bytes=size_bytes,
            last_modified=last_modified,
            fingerprint=fingerprint,
        )

        return [Document(
            uuid=str(uuid4()),
            content=content,
            lang=lang,
            tags=[],
            source=source_config,
            metadata=meta_config,
            chunk_groups=[]
        )]

    async def _caption_image(self, path: str) -> str:
        if not self.captioner:
            logger.warning("Captioner not provided; returning empty string.", extra={"file": path})
            return ""
        try:
            async with log_duration("caption_image", path=path):
                image_bytes = await load_image_bytes_from_source(path)
                captions = await self.captioner.generate_captions([image_bytes])
                return captions[0] if captions else ""
        except Exception as e:
            logger.warning(f"Failed to caption image {path}: {e}")
            return ""

    async def _transcribe_audio(self, path: str, mime: str) -> str:
        if not self.transcriber:
            logger.warning("Transcriber not provided; returning empty string.", extra={"file": path})
            return ""
        try:
            async with log_duration("transcribe_audio", path=path):
                audio_bytes = await load_file_from_path(path)
                return await self.transcriber.transcribe(audio_bytes, mime)
        except Exception as e:
            logger.exception("Failed to transcribe audio", extra={"file": path, "error": str(e)})
            return ""

    async def _read_text(self, path: str) -> str:
        async with log_duration("read_text", path=path):
            def read():
                with open(path, 'rb') as f:
                    raw_start = f.read(2048)
                encoding = chardet.detect(raw_start).get("encoding", "utf-8")
                with open(path, encoding=encoding, errors="replace") as f:
                    return f.read()
            return await asyncio.to_thread(read)

    async def _read_html(self, path: str) -> str:
        try:
            from markdownify import markdownify as md
        except ImportError:
            raise ImportError("markdownify is required to convert HTML.")
        html = await self._read_text(path)
        async with log_duration("read_html", path=path):
            return await asyncio.to_thread(md, html)

    async def _read_pdf(self, path: str) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required to read PDF files.")
        async with log_duration("read_pdf", path=path):
            def extract_pdf():
                reader = PdfReader(path)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            return await asyncio.to_thread(extract_pdf)

    async def _read_docx(self, path: str) -> str:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required to read DOCX files.")
        async with log_duration("read_docx", path=path):
            def extract_docx():
                doc = docx.Document(path)
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return await asyncio.to_thread(extract_docx)

    async def _read_json(self, path: str) -> str:
        async with log_duration("read_json", path=path):
            def parse_json():
                with open(path, 'rb') as f:
                    raw = f.read()
                encoding = chardet.detect(raw).get("encoding", "utf-8")
                data = json.loads(raw.decode(encoding, errors="replace"))
                return json.dumps(data, separators=(",", ":"), ensure_ascii=False)
            return await asyncio.to_thread(parse_json)

    async def _hash_file(self, path: Path) -> str:
        def hash_file():
            hasher = hashlib.sha256()
            with path.open("rb") as f:
                while chunk := f.read(8192):
                    hasher.update(chunk)
            return hasher.hexdigest()

        return await asyncio.to_thread(hash_file)
